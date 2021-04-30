from docx.enum.style import WD_STYLE_TYPE
import docx
import pymysql
import pytesseract
import re
import os 

docname = 'C:\\Users\\Duke Maquiling\\Documents\\10.10.100.50 Documentation\\Documentations\\global1.docx'
tess = 'C:\\Users\\Duke Maquiling\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract'
rowpng = 'C:\\Users\\Duke Maquiling\\Documents\\10.10.100.50 Documentation\\Images\\_global1rows.png'
tablepng = 'C:\\Users\\Duke Maquiling\\Documents\\10.10.100.50 Documentation\\Images\\_global1table.png'
updatedpng = 'C:\\Users\\Duke Maquiling\\Documents\\10.10.100.50 Documentation\\Images\\testifwupdated.png'
sizepng = 'C:\\Users\\Duke Maquiling\\Documents\\10.10.100.50 Documentation\\Images\\_global1size.png'

######################################################################FUNCTIONS########################################################################################


def delete_paragraph(paragraph):
    #Deletes .add_paragraph sections 
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def tesseract(s):
    #Sets up the tesseract package
    #Ensure to change directories respective to your own respective directory
    pytesseract.pytesseract.tesseract_cmd = s

def size2list(s):
    #Takes an image of all the sizes and parses into text. Stores into an list
    #Ensure to change directories to your own respective directory
    sizes = [] 
    sizeParse = pytesseract.image_to_string(s)
    #Size image to text parsing 
    res = [sizeParse[ele.start():ele.end()] for ele in re.finditer(r'\S+', sizeParse)]
    i = 0
    while len(res) > i:
        if "KiB" in res[i]: #i dont like how this is hardcoded 
            s = res[i]
            #print(s[0:3] + " " + s[3:6])
            sizes.append(s[0:3] + " " + s[3:6])
            i+=1
        elif "MiB" in res[i]: #i dont like how this is hardcoded 
            s = res[i]
            #print(s[0:3] + " " + s[3:6])
            sizes.append(s[0:3] + " " + s[3:6])
            i+=1
        elif "KB" in res[i]: #i dont like how this is hardcoded 
            s = res[i]
            #print(s[0:2] + " " + s[2:4])
            sizes.append(s[0:3] + " " + s[3:6])
            i+=1
        elif "MB" in res[i]: #i dont like how this is hardcoded 
            s = res[i]
            #print(s[0:3] + " " + s[3:6])
            sizes.append(s[0:3] + " " + s[3:6])
            i+=1
        else:
            #print(res[i]+" "+res[i+1])
            sizes.append(res[i]+" "+res[i+1])
            i+=2
    return sizes

def update2list(s):
    #Takes an image of all last updated dates and parses into text. Stores parsed text into a list
    #Ensure to change directories to your own respective directory
    lastupdated = ['N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A']
    
    lastUpdateParse = pytesseract.image_to_string(s)
    res1 = [lastUpdateParse[ele.start():ele.end()] for ele in re.finditer(r'\S+', lastUpdateParse)]
    k=0

    while len(res1) > k:
            if 'N/A' in res1[k]:
                
                #print(str(k) + " "+ res1[k])
                #lastupdated.append("N/A")
                k+=1
            elif "Nia" in res1[k]:

                #print(res1[k])
                #lastupdated.append("N/A")
                #print(k)
                k+=1
            elif "N/a" in res1[k]:
   
                #print(res1[k])
                #lastupdated.append("N/A")
                #print(k)
                k+=1
            elif "NA" in res1[k]:
                
                #print(res1[k])
                #lastupdated.append("N/A")
                #print(k)
                k+=1
            else: 
                #print(str(k) + " " + res1[k]+" "+res1[k+1]+" "+res1[k+2]+" "+res1[k+3]+" "+res1[k+4]+" "+res1[k+5] + "',")
                #lastupdated.append(res1[k]+" "+res1[k+1]+" "+res1[k+2]+" "+res1[k+3]+" "+res1[k+4]+" "+res1[k+5])
                k+=6
    
    return lastupdated

def row2list(s):
    rows = ['0', '3', '0', '170', '3', '4', '0', '1', '1', '1', '18', '1']
    rowsparse = pytesseract.image_to_string(s)
    res2 = [rowsparse[ele.start():ele.end()] for ele in re.finditer(r'\S+', rowsparse)]
    
    #k=0
    #while len(res2) > k:
    #    rows.append(res2[k])
    #    k+=1

    return rows

def table2list(s):

    name = []
    nameParse = pytesseract.image_to_string(s)
    res = [nameParse[ele.start():ele.end()] for ele in re.finditer(r'\S+', nameParse)]
    
    #for n in res:
    #    print(n)

    return res
    

def readdocx(s):
    #Reads into the document for docx
    #Ensure to change directories to your own respective directory
    return docx.Document(s) 

def savedocx(s):
    #Saves the word docx 
    #Ensure to change directories to your own respective directory
    try:
        doc.save(s)
        print("Finish formatting the document.")
    except PermissionError:
        print("SCRIPT TERMINATED: Please close the word Document that you are trying to open and run the script again.\nREASON: Document will open, but no changes will occur")
        exit()

def opendocx(s):
    #Automatically opens up the word doc in question 
    #Ensure to change directories to your own respective directory
    os.startfile(s)
    print("Openning Word Document right now...")



##################################################################################################################################################################
    
#Need to fix this. Make sure it's taking a picture and parsing it 


#name = "\\WLRdb.docx"
tesseract(tess)
sizes = size2list(sizepng)
print("Sizes: " + str(len(sizes)))
lastupdated = update2list(updatedpng)
print("Last Updated: " + str(len(lastupdated)))
rows = row2list(rowpng)
print("Rows: " + str(len(rows)))
names = table2list(tablepng)
print("Names: " + str(len(names)))

#Word Document Font Styling 
doc = readdocx(docname) 
styles = doc.styles
try:
    style = styles.add_style('heading', WD_STYLE_TYPE.PARAGRAPH)
except ValueError:
    print('headings style is already included into the word doc')

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = docx.shared.Pt(26)

p = doc.add_paragraph()
p.add_run("10.10.100.50").bold=True
p.alignment = 1
p = doc.add_paragraph()
p.add_run("SQL Tables").bold=True
p.alignment = 1

style = doc.styles['heading']
font = style.font
font.name = 'Calibri'
font.size = docx.shared.Pt(18)
p = doc.add_paragraph()
p.add_run("Database: RGL")
p.alignment = 1
doc.add_page_break()

j = 0
while len(names) > j: 
    if rows[j] == "0":
        p = doc.add_paragraph()
        p.add_run("TABLE: ").bold=True
        p.add_run(names[j])
        p = doc.add_paragraph()
        p.add_run("SIZE: ").bold=True
        p.add_run(sizes[j])
        p = doc.add_paragraph()
        p.add_run("OLDEST RECORD: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("LAST UPDATED: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("UPDATED FROM: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("PROGRAM: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("ACCESS FROM WHERE: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("NUMBER OF ROWS: ").bold=True
        p.add_run(rows[j])
        doc.add_page_break()
        doc.add_paragraph(" ")
    else:
        p = doc.add_paragraph()
        p.add_run("TABLE: ").bold=True
        p.add_run(names[j])
        p = doc.add_paragraph()
        p.add_run("SIZE: ").bold=True
        p.add_run(sizes[j])
        p = doc.add_paragraph()
        p.add_run("OLDEST RECORD: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("LAST UPDATED: ").bold=True
        p.add_run(lastupdated[j])
        p = doc.add_paragraph()
        p.add_run("UPDATED FROM: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("PROGRAM: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("ACCESS FROM WHERE: ").bold=True
        p.add_run("N/A")
        p = doc.add_paragraph()
        p.add_run("NUMBER OF ROWS: ").bold=True
        p.add_run(rows[j])
        doc.add_paragraph(" ")
        doc.add_page_break()
    j+=1

savedocx(docname)
opendocx(docname)






